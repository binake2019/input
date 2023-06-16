from docx import Document
import time

#打开word文档
document = Document("常涛.docx")

##获取所有段落
##all_paragraphs = document.paragraphs
##打印看看all_paragraphs是什么东西
##print(type(all_paragraphs)) #<class 'list'>，打印后发现是列表
##是列表就开始循环读取
##for paragraph in all_paragraphs:
##    
##    #打印每一个段落的文字
##    print(paragraph.text)
##
table = document.tables[0]


table_temp = []

ls = list()
for row in table.rows:

    temp = list()
    for cell in row.cells:

        if cell.text not in temp:
            temp.append(cell.text)
            print(cell.text,end=' ')
    print('\n')
    print('end')
    ls.append(temp)

print(ls)








# 获取word信息，1姓名，2性别，3文化程度，4工作单位，5单位地址，6通讯地址，
# 7邮编，8电话，9申请作业项目（叉车司机），10简历 ，11，同意
lsinfo = []
##
##x,y = (0,1)
##d = table.cell(x,y).text
##if d == "姓名":
##    d = table.cell(x,y+1).text
##lsinfo.append(d)
##
##x,y = (0,2)
##d = table.cell(x,y).text
##if d == "性别":
##    d = table.cell(x,y+1).text
##lsinfo.append(d)
##    
##x,y = (1,1)
##d = table.cell(x,y).text
##if d == "身份证件号":
##    d = table.cell(x,y+1).text
##lsinfo.append(d)
##
##x,y = (1,2)
##d = table.cell(x,y).text
##if d == "文化程度":
##    d = table.cell(x,y+1).text
##lsinfo.append(d)
##
##
##print(lsinfo)
